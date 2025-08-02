from flask import Blueprint, request, jsonify, send_file
from werkzeug.utils import secure_filename
import os
import re
from docx import Document
import phonenumbers
from phonenumbers import geocoder, carrier
import tempfile
import uuid
from datetime import datetime

redaction_bp = Blueprint('redaction', __name__)

# Configure upload settings
UPLOAD_FOLDER = '/tmp/uploads'
PROCESSED_FOLDER = '/tmp/processed'
ALLOWED_EXTENSIONS = {'txt', 'docx'}  # Removed PDF for deployment

# Ensure directories exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PROCESSED_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

class PIIDetector:
    def __init__(self):
        # Email regex pattern
        self.email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
        
        # SSN patterns (US format)
        self.ssn_pattern = re.compile(r'\b\d{3}-?\d{2}-?\d{4}\b')
        
        # Credit card patterns (basic)
        self.credit_card_pattern = re.compile(r'\b(?:\d{4}[-\s]?){3}\d{4}\b')
        
        # Date patterns
        self.date_pattern = re.compile(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b')
        
        # IBAN pattern (basic)
        self.iban_pattern = re.compile(r'\b[A-Z]{2}\d{2}[A-Z0-9]{4}\d{7}([A-Z0-9]?){0,16}\b')
        
        # Address pattern (basic)
        self.address_pattern = re.compile(r'\b\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Lane|Ln|Drive|Dr)\b', re.IGNORECASE)

    def detect_pii(self, text):
        """Detect PII in text and return list of matches with positions"""
        pii_items = []
        
        # Email detection
        for match in self.email_pattern.finditer(text):
            pii_items.append({
                'type': 'email',
                'text': match.group(),
                'start': match.start(),
                'end': match.end(),
                'confidence': 'high'
            })
        
        # Phone number detection
        try:
            for match in phonenumbers.PhoneNumberMatcher(text, None):
                pii_items.append({
                    'type': 'phone',
                    'text': text[match.start:match.end],
                    'start': match.start,
                    'end': match.end,
                    'confidence': 'high'
                })
        except:
            pass
        
        # SSN detection
        for match in self.ssn_pattern.finditer(text):
            pii_items.append({
                'type': 'ssn',
                'text': match.group(),
                'start': match.start(),
                'end': match.end(),
                'confidence': 'high'
            })
        
        # Credit card detection
        for match in self.credit_card_pattern.finditer(text):
            pii_items.append({
                'type': 'credit_card',
                'text': match.group(),
                'start': match.start(),
                'end': match.end(),
                'confidence': 'medium'
            })
        
        # Date detection
        for match in self.date_pattern.finditer(text):
            pii_items.append({
                'type': 'date',
                'text': match.group(),
                'start': match.start(),
                'end': match.end(),
                'confidence': 'low'
            })
        
        # IBAN detection
        for match in self.iban_pattern.finditer(text):
            pii_items.append({
                'type': 'iban',
                'text': match.group(),
                'start': match.start(),
                'end': match.end(),
                'confidence': 'high'
            })
        
        # Address detection
        for match in self.address_pattern.finditer(text):
            pii_items.append({
                'type': 'address',
                'text': match.group(),
                'start': match.start(),
                'end': match.end(),
                'confidence': 'medium'
            })
        
        return sorted(pii_items, key=lambda x: x['start'])

class DocumentProcessor:
    def __init__(self):
        self.pii_detector = PIIDetector()
    
    def process_docx(self, file_path):
        """Process DOCX file and detect PII"""
        doc = Document(file_path)
        full_text = []
        pii_items = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        text_content = '\n'.join(full_text)
        pii_items = self.pii_detector.detect_pii(text_content)
        
        return {
            'text': text_content,
            'pii_items': pii_items,
            'paragraphs': full_text
        }
    
    def process_txt(self, file_path):
        """Process TXT file and detect PII"""
        with open(file_path, 'r', encoding='utf-8') as file:
            text_content = file.read()
        
        pii_items = self.pii_detector.detect_pii(text_content)
        
        return {
            'text': text_content,
            'pii_items': pii_items
        }
    
    def redact_docx(self, file_path, redaction_items, output_path):
        """Apply redactions to DOCX file"""
        doc = Document(file_path)
        
        for para in doc.paragraphs:
            for item in redaction_items:
                if item['text'] in para.text:
                    para.text = para.text.replace(item['text'], '[REDACTED]')
        
        doc.save(output_path)
        return output_path
    
    def redact_txt(self, file_path, redaction_items, output_path):
        """Apply redactions to TXT file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        for item in redaction_items:
            content = content.replace(item['text'], '[REDACTED]')
        
        with open(output_path, 'w', encoding='utf-8') as file:
            file.write(content)
        
        return output_path

# Initialize processor
processor = DocumentProcessor()

@redaction_bp.route('/upload', methods=['POST'])
def upload_file():
    """Upload and analyze document for PII"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': 'File type not supported. Please use TXT or DOCX files.'}), 400
    
    try:
        # Generate unique filename
        file_id = str(uuid.uuid4())
        filename = secure_filename(file.filename)
        file_extension = filename.rsplit('.', 1)[1].lower()
        stored_filename = f"{file_id}.{file_extension}"
        file_path = os.path.join(UPLOAD_FOLDER, stored_filename)
        
        # Save uploaded file
        file.save(file_path)
        
        # Process file based on type
        if file_extension == 'docx':
            result = processor.process_docx(file_path)
            document_type = 'docx'
        elif file_extension == 'txt':
            result = processor.process_txt(file_path)
            document_type = 'txt'
        
        return jsonify({
            'file_id': file_id,
            'filename': filename,
            'document_type': document_type,
            'analysis': result,
            'upload_time': datetime.now().isoformat()
        })
    
    except Exception as e:
        return jsonify({'error': f'Processing failed: {str(e)}'}), 500

@redaction_bp.route('/redact', methods=['POST'])
def redact_document():
    """Apply redactions to document"""
    data = request.get_json()
    
    if not data or 'file_id' not in data or 'redactions' not in data:
        return jsonify({'error': 'Missing required data'}), 400
    
    file_id = data['file_id']
    redactions = data['redactions']
    document_type = data.get('document_type', 'txt')
    
    try:
        # Find original file
        original_files = [f for f in os.listdir(UPLOAD_FOLDER) if f.startswith(file_id)]
        if not original_files:
            return jsonify({'error': 'Original file not found'}), 404
        
        original_file = original_files[0]
        original_path = os.path.join(UPLOAD_FOLDER, original_file)
        
        # Generate output filename
        output_filename = f"{file_id}_redacted.{document_type}"
        output_path = os.path.join(PROCESSED_FOLDER, output_filename)
        
        # Apply redactions based on document type
        if document_type == 'docx':
            result_path = processor.redact_docx(original_path, redactions, output_path)
        elif document_type == 'txt':
            result_path = processor.redact_txt(original_path, redactions, output_path)
        
        return jsonify({
            'success': True,
            'output_file': output_filename,
            'download_url': f'/api/redaction/download/{output_filename}'
        })
    
    except Exception as e:
        return jsonify({'error': f'Redaction failed: {str(e)}'}), 500

@redaction_bp.route('/download/<filename>')
def download_file(filename):
    """Download processed file"""
    try:
        return send_file(
            os.path.join(PROCESSED_FOLDER, filename),
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        return jsonify({'error': f'Download failed: {str(e)}'}), 404

@redaction_bp.route('/health')
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'supported_formats': list(ALLOWED_EXTENSIONS),
        'note': 'PDF support available in local version'
    })

